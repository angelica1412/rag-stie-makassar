# Pembacaan Dokumen

import os
import re
import fitz
import pdfplumber
from llama_index.core import Document

from docx import Document as DocxDocument

import openpyxl

NARATIF_PATH = "./data/documents/naratif"
FORM_PATH = "./data/documents/form"


def decode_filename(filename: str) -> str:
    """Decode karakter URL-encoded dalam nama file."""
    # _20 = spasi (%20), _2F = slash (%2F), dll
    decoded = filename
    decoded = decoded.replace("_2F", "/")
    decoded = decoded.replace("_2B", "+")
    decoded = decoded.replace("_20", " ")
    return decoded


def clean_text(text: str) -> str:
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_tables_from_page(plumber_page) -> str:
    table_text = ""
    try:
        tables = plumber_page.extract_tables()
        for table in tables:
            for row in table:
                cleaned = [cell.strip() if cell else "" for cell in row]
                table_text += " | ".join(cleaned) + "\n"
    except Exception:
        pass
    return table_text


# ── Ekstraksi metadata per tipe file ─────────────────────────────────────────

def extract_metadata_from_pdf(filepath: str) -> str:
    """Ambil teks dari halaman pertama PDF."""
    try:
        fitz_doc = fitz.open(filepath)
        text = fitz_doc[0].get_text("text")
        fitz_doc.close()
        text = clean_text(text)
        return text[:500] if len(text) > 500 else text
    except Exception as e:
        return f"Tidak dapat membaca PDF: {e}"


def extract_metadata_from_docx(filepath: str) -> str:
    """Ambil teks dari paragraf pertama file docx."""
    try:
        doc = DocxDocument(filepath)
        texts = []
        for para in doc.paragraphs[:20]:  # ambil 20 paragraf pertama
            if para.text.strip():
                texts.append(para.text.strip())
        result = "\n".join(texts)
        return result[:500] if len(result) > 500 else result
    except Exception as e:
        return f"Tidak dapat membaca DOCX: {e}"


def extract_metadata_from_xlsx(filepath: str) -> str:
    """Ambil teks dari baris pertama file xlsx."""
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        ws = wb.active
        texts = []
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            if row_count >= 10:  # ambil 10 baris pertama
                break
            row_text = " | ".join(
                str(cell) for cell in row if cell is not None
            )
            if row_text.strip():
                texts.append(row_text)
                row_count += 1
        wb.close()
        result = "\n".join(texts)
        return result[:500] if len(result) > 500 else result
    except Exception as e:
        return f"Tidak dapat membaca XLSX: {e}"


# ── Fungsi utama baca dokumen ─────────────────────────────────────────────────

def read_naratif_documents() -> list[Document]:
    """
    Baca dokumen naratif (pedoman, standar, manual, IK) secara full-text.
    Hanya membaca file PDF.
    """
    documents = []
    print("\n=== Membaca dokumen NARATIF ===")

    if not os.path.exists(NARATIF_PATH):
        print(f"Folder tidak ditemukan: {NARATIF_PATH}")
        return documents

    for filename in os.listdir(NARATIF_PATH):
        if not filename.endswith(".pdf"):
            continue

        filepath = os.path.join(NARATIF_PATH, filename)
        display_name = decode_filename(filename)
        print(f"Membaca: {display_name}")

        try:
            fitz_doc = fitz.open(filepath)
            plumber_doc = pdfplumber.open(filepath)
            page_count = 0

            for page_num in range(len(fitz_doc)):
                fitz_page = fitz_doc[page_num]
                text = fitz_page.get_text("text")
                text = clean_text(text)

                plumber_page = plumber_doc.pages[page_num]
                table_text = extract_tables_from_page(plumber_page)

                page_content = text
                if table_text:
                    page_content += "\n[TABEL]\n" + table_text

                if page_content.strip():
                    doc = Document(
                        text=page_content,
                        metadata={
                            "file_name": decode_filename(filename),
                            "file_path": filepath,
                            "source": decode_filename(filename),
                            "page_number": page_num + 1,
                            "tipe_dokumen": "naratif"
                        }
                    )
                    documents.append(doc)
                    page_count += 1

            fitz_doc.close()
            plumber_doc.close()
            print(f"  Berhasil: {page_count} halaman dari {display_name}")

        except Exception as e:
            print(f"  Error membaca {filename}: {e}")

    # Ringkasan
    naratif_files = set(doc.metadata.get("file_name", "") for doc in documents)
    print(f"Total naratif: {len(naratif_files)} file ({len(documents)} halaman)")
    return documents


def read_form_documents() -> list[Document]:
    """
    Baca dokumen form — metadata only.
    Mendukung: PDF, DOCX, DOC, XLSX.
    """
    documents = []
    print("\n=== Membaca dokumen FORM (metadata only) ===")

    if not os.path.exists(FORM_PATH):
        print(f"Folder tidak ditemukan: {FORM_PATH}")
        return documents

    # Format yang didukung
    supported_ext = (".pdf", ".docx", ".doc", ".xlsx")

    for filename in os.listdir(FORM_PATH):
        ext = os.path.splitext(filename)[1].lower()
        if ext not in supported_ext:
            continue

        filepath = os.path.join(FORM_PATH, filename)
        decoded_name = decode_filename(filename)
        print(f"Membaca form: {decoded_name}")

        try:
            # Ekstrak metadata sesuai tipe file
            if ext == ".pdf":
                metadata_text = extract_metadata_from_pdf(filepath)
                tipe_file = "PDF"
            elif ext in (".docx", ".doc"):
                metadata_text = extract_metadata_from_docx(filepath)
                tipe_file = "Word Document"
            elif ext == ".xlsx":
                metadata_text = extract_metadata_from_xlsx(filepath)
                tipe_file = "Excel Spreadsheet"
            else:
                continue

            # Buat deskripsi dokumen form
            doc_content = (
                f"Formulir: {decoded_name}\n"
                f"Tipe File: {tipe_file}\n"
                f"Tipe Dokumen: Formulir\n"
                f"Deskripsi: {metadata_text}\n"
                f"Kegunaan: Formulir ini tersedia untuk digunakan "
                f"sesuai kebutuhan akademik dan administratif kampus."
            )

            if doc_content.strip():
                doc = Document(
                    text=doc_content,
                    metadata={
                        "file_name": decoded_name,
                        "file_path": filepath,
                        "source": decoded_name,
                        "page_number": 1,
                        "tipe_dokumen": "form",
                        "tipe_file": tipe_file
                    }
                )
                documents.append(doc)
                print(f"  Berhasil: {tipe_file} — {decoded_name}")

        except Exception as e:
            print(f"  Error membaca {filename}: {e}")

    print(f"Total form: {len(documents)} file")
    return documents


def read_pdfs_from_folder(folder_path: str = None) -> list[Document]:
    """
    Fungsi utama — baca semua dokumen dari kedua subfolder.
    - Naratif: full text + tabel (PDF only)
    - Form: metadata only (PDF, DOCX, DOC, XLSX)
    """
    all_documents = []

    naratif_docs = read_naratif_documents()
    all_documents.extend(naratif_docs)

    form_docs = read_form_documents()
    all_documents.extend(form_docs)

    print(f"\n=== Ringkasan Ingestion ===")
    naratif_files = set(
        doc.metadata.get("file_name", "")
        for doc in naratif_docs
    )
    print(f"Dokumen naratif : {len(naratif_files)} file "
          f"({len(naratif_docs)} halaman)")
    print(f"Dokumen form    : {len(form_docs)} file")
    print(f"Total chunks    : {len(all_documents)}")

    return all_documents